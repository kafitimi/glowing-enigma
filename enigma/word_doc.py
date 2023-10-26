"""
Функции работы с документами
"""
import sys
from os.path import join

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docxtpl import DocxTemplate

CENTER = 'Table Heading'
JUSTIFY = 'Table Contents'


def get_template(filename: str) -> DocxTemplate:
    """ Читаем шаблон РПД """
    try:
        template = DocxTemplate(join('templates', filename))
    except OSError:
        print('Не могу открыть шаблон')
        sys.exit()
    return template


def set_cell_text(table: Table, row: int, col: int, style: str, text: str) -> None:
    """ Добавить текст в ячейку таблицы """
    cell = table.cell(row, col)
    if cell.text:
        cell.add_paragraph(text, style)
    else:
        cell.text = text
        cell.paragraphs[0].style = style


def add_table_rows(table: Table, rows: int) -> None:
    """ Добавить строки в таблицу """
    def get_border(name):
        border = OxmlElement(f'w:{name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        return border

    def get_borders():
        borders = OxmlElement("w:tcBorders")
        borders.append(get_border('top'))
        borders.append(get_border('left'))
        borders.append(get_border('bottom'))
        borders.append(get_border('right'))
        return borders

    for _ in range(rows):
        row = table.add_row()
        for cell in row.cells:
            props = cell._element.tcPr  # pylint: disable=protected-access
            props.append(get_borders())


def remove_table(template: DocxTemplate, table_index: int) -> None:
    """ Удаляем таблицу из шаблона по её индексу """
    docx = template.get_docx()
    table_element = docx.tables[table_index]._element  # pylint: disable=protected-access
    parent_element = table_element.getparent()
    parent_element.remove(table_element)
