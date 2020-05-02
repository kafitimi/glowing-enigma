"""
Гвардейский краснознаменный корчеватель РПД имени Э. И. Шамаева
"""
import os

from docx import Document


def findout_fos(documento):
    """ Извлекаем текст раздела "Фонд оценочных средств" из РПД """
    text_at_start = ['Фонд оценочных средств']
    text_at_end = ['учебной литературы', 'iprbook', 'lanbook', 'НБ СВФУ',
                   'информационно-телекоммуникационной сети']
    flag, result_text = False, ''
    for paragraph in documento.paragraphs:
        if text_at_start[0] in paragraph.text:
            flag = True
        if any(key_word in paragraph.text for key_word in text_at_end):
            flag = False
        if flag:
            result_text += (paragraph.text + '\n')
    result_text = result_text.replace('\n\n', '\n').replace('  ', ' ', 1000)
    return result_text


def findout_fos_table(documento):
    """ Извлекаем результаты обучения (ЗУВы) по дисциплине из РПД """
    text_in_table = ['Коды оцениваемых компетенций', 'Показатель оценивания',
                     'Шкалы оценивания', 'Уровни освоения', 'Критерии оценивания',
                     'оцениваемых компетенций', 'показатель оценивания',
                     'шкалы оценивания', 'уровни освоения', 'критерии оценивания',
                     'тлично', 'хорошо', 'удовл', 'зачтено']
    tablez = ''
    for table in documento.tables:
        flag, tablelist = False, []
        max_col_num = 0
        for row in table.rows:
            rowlist = []
            for cell in row.cells:
                celltext = ''
                for paragraph in cell.paragraphs:
                    celltext += ('\n' + paragraph.text)
                rowlist.append(celltext)
            if len(rowlist) > max_col_num:
                max_col_num = len(rowlist)
            if rowlist:
                tablelist.append(rowlist)
        if max_col_num > 0:
            tablelist2 = []
            for row in tablelist:
                if len(row) < max_col_num:
                    tablelist2.append(row.extend([' '] * (max_col_num - len(row))))
                else:
                    if row:
                        tablelist2.append(row)
            # print(len(row), max_col_num, tablelist2)
            # print(tablelist2[-1] is not None)
            tablelist = [list(line) for line in zip(*tablelist2)]
            tabletext2 = '\n'.join(['\n'.join(list(dict.fromkeys(line))) for line in tablelist])
            if any(key_word in tabletext2 for key_word in text_in_table):
                tablez += tabletext2.replace('\n\n', '\n').replace('  ', ' ', 1000)
    return tablez


for filename in os.listdir():
    if filename.endswith(".docx"):
        f = open(filename, 'rb')
        document = Document(f)
        got_fos = findout_fos_table(document) + findout_fos(document)
        f.close()
        f2 = open(filename[:-4] + 'txt', "w+")
        f2.write(got_fos)
        print(filename)
        print(findout_fos_table(document))
