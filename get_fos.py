""" Генерация ФОС """
import os
from argparse import ArgumentParser, Namespace
from typing import Dict, List

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import CT_P, CT_Tbl
from docx.shared import Cm
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate

import core


def iterate_items(parent):
    """ Обход параграфов и таблиц в документе """
    if isinstance(parent, DocumentType):
        parent_elem = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elem = parent._tc  # pylint: disable=protected-access
    else:
        raise ValueError('Oops')

    for child in parent_elem.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def get_section_paragraphs(input_filename: str, start_kw: List[str], final_kw: List[str]) -> List[str]:
    """ Извлечь список абзацев текста из docx-файла """
    source = Document(input_filename)
    started = False
    paragraphs = []
    for item in iterate_items(source):
        if not started:
            if isinstance(item, Paragraph) and any(kw in item.text for kw in start_kw):
                started = True
        else:
            if isinstance(item, Paragraph) and any(kw in item.text for kw in final_kw):
                break
            if isinstance(item, Paragraph):
                text = item.text.strip()
                if text:
                    paragraphs.append(text + '\n')
            elif isinstance(item, Table):
                pass
    return paragraphs


def get_rpd(name):
    """ Получить список имен файлов РПД """
    result, base_dir = None, 'rpds'
    for file_name in os.listdir(base_dir):
        if file_name.endswith('.docx'):
            if name in file_name:
                result = os.path.join(base_dir, file_name)
                break
    return result


def fill_table_1(template: DocxTemplate, context: Dict[str, any]) -> None:
    """ Заполнение таблиц с формами контроля """
    # control_fancy_name = {
    #     core.CT_EXAM: 'Экзамен',
    #     core.CT_CREDIT_GRADE: 'Зачет с оценкой',
    #     core.CT_CREDIT: 'Зачет',
    #     core.CT_COURSEWORK: 'Курсовой проект',
    # }

    plan: core.EducationPlan = context['plan']
    if plan.degree == core.BACHELOR:
        core.remove_table(template, 1)
    elif plan.degree == core.MASTER:
        core.remove_table(template, 2)
    table: Table = template.get_docx().tables[1]

    row_number = 0
    for competence in sorted(plan.competence_codes.values(), key=core.Competence.repr):
        core.add_table_rows(table, 1)
        row = len(table.rows) - 1
        row_number += 1
        core.set_cell_text(table, row, 0, core.CENTER, str(row_number))
        core.set_cell_text(table, row, 1, core.JUSTIFY, competence.code + ' ' + competence.description)
        for number in range(1, 9 if plan.degree == core.BACHELOR else 5):
            cell_components = []
            subjects = [plan.subject_codes[s] for s in competence.subjects]
            subjects.sort(key=lambda s: s.code)
            for subject in subjects:
                for sem_num, sem_work in subject.semesters.items():
                    if sem_num == number:
                        if sem_work.control:
                            values = subject.code, subject.name, ', '.join(sem_work.control)
                            cell_components.append('%s %s (%s)' % values)
                        else:
                            values = subject.code, subject.name
                            cell_components.append('%s %s' % values)
            core.set_cell_text(table, row, number + 1, core.CENTER, '\n'.join(cell_components))
    core.fix_table_borders(table)

    rpd_dict = context['rpd_dict']
    table: Table = template.get_docx().tables[2]
    row_number1 = 0
    for competence in sorted(plan.competence_codes.values(), key=core.Competence.repr):
        core.add_table_rows(table, 1)
        row = len(table.rows) - 1
        row_number1 += 1
        core.set_cell_text(table, row, 0, core.CENTER, str(row_number1))
        core.set_cell_text(table, row, 1, core.JUSTIFY, competence.code + ' ' + competence.description)
        values = ['%s %s' % (k, v.description) for k, v in competence.indicator_codes.items()]
        core.set_cell_text(table, row, 2, core.JUSTIFY, '\n'.join(values))
        row_number2 = 0
        subjects = [plan.subject_codes[s] for s in competence.subjects]
        for subject in subjects:
            if not subject.code.startswith('Б1'):
                continue
            zuv_raw = ''
            rpd = rpd_dict.get(subject.code)
            if rpd:
                zuv = rpd.competences.get(competence.code)
                if zuv:
                    zuv_raw = zuv.raw
            core.add_table_rows(table, 1)
            row = len(table.rows) - 1
            row_number2 += 1
            core.set_cell_text(table, row, 0, core.CENTER, '%d.%d' % (row_number1, row_number2))
            core.set_cell_text(table, row, 1, core.JUSTIFY, subject.code + ' ' + subject.name)
            core.set_cell_text(table, row, 3, core.JUSTIFY, zuv_raw)
            core.set_cell_text(table, row, 4, core.JUSTIFY, zuv_raw)
            core.set_cell_text(table, row, 5, core.JUSTIFY, zuv_raw)
            core.set_cell_text(table, row, 6, core.JUSTIFY, zuv_raw)
    core.fix_table_borders(table)


def fill_table_2_1(template: DocxTemplate, context: Dict[str, any]) -> None:
    """ Заполнение таблицы в разделе 2.1 """
    plan: core.EducationPlan = context['plan']
    table: Table = template.get_docx().tables[3]
    for subject in sorted(plan.subject_codes.values(), key=core.Subject.repr):
        core.add_table_rows(table, 1)
        row_index = len(table.rows) - 1
        core.set_cell_text(table, row_index, 0, core.CENTER, subject.code)
        core.set_cell_text(table, row_index, 1, core.JUSTIFY, subject.name)
    core.fix_table_borders(table)


def fill_section_2_2(template: DocxTemplate, context: Dict[str, any]) -> None:
    """ Заполнение раздела 2.2 """
    marker = None
    for paragraph in template.get_docx().paragraphs:
        keywords = ['оценочные средства для', 'государственной итоговой аттестации']
        if all(kw in paragraph.text.lower() for kw in keywords):
            marker = paragraph
            break

    plan: core.EducationPlan = context['plan']
    subjects = sorted(plan.subject_codes.values(), key=core.Subject.repr)
    for subj in subjects:
        rpd = get_rpd(subj.name)
        if not rpd:
            continue

        paragraph = marker.insert_paragraph_before('%s %s' % (subj.code, subj.name))
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        paragraph.paragraph_format.first_line_indent = Cm(0)
        for run in paragraph.runs:
            run.bold = True

        for item in iterate_items(Document(rpd)):
            if isinstance(item, Paragraph):
                paragraph = marker.insert_paragraph_before(item.text)
                paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                paragraph.paragraph_format.first_line_indent = Cm(0)


def fill_table_4(template: DocxTemplate, context: Dict[str, any]) -> None:
    """ Заполнение бланка "Лист сформированности компетенций" """
    plan: core.EducationPlan = context['plan']
    table: Table = template.get_docx().tables[-1]
    row_number = 0
    for competence in sorted(plan.competence_codes.values(), key=core.Competence.repr):
        core.add_table_rows(table, 1)
        row_index = len(table.rows) - 1
        row_number += 1
        core.set_cell_text(table, row_index, 0, core.CENTER, str(row_number))
        core.set_cell_text(table, row_index, 1, core.JUSTIFY, competence.code + ' ' + competence.description)
        subjects = [plan.subject_codes[s] for s in competence.subjects]
        for subject in sorted(subjects, key=core.Subject.repr):
            core.add_table_rows(table, 1)
            row_index = len(table.rows) - 1
            core.set_cell_text(table, row_index, 1, core.JUSTIFY, subject.code + ' ' + subject.name)

    core.add_table_rows(table, 1)
    row_number += 1
    row_index = len(table.rows) - 1
    core.set_cell_text(table, row_index, 0, core.CENTER, str(row_number))
    core.set_cell_text(table, row_index, 1, core.JUSTIFY, 'Практики')

    core.add_table_rows(table, 1)
    row_number += 1
    row_index = len(table.rows) - 1
    core.set_cell_text(table, row_index, 0, core.CENTER, str(row_number))
    core.set_cell_text(table, row_index, 1, core.JUSTIFY, 'НИР')

    core.fix_table_borders(table)


def get_rpd_dict(plan: core.EducationPlan, rpd_dir: str) -> Dict[str, core.RPD]:
    result = {}
    for filename in os.listdir(rpd_dir):
        if filename.startswith('~'):
            continue
        if not filename.endswith('.docx'):
            continue
        code, name = tuple(filename.split(' ', 1))
        full_path = os.path.join(rpd_dir, filename)
        rpd_doc = Document(full_path)
        comp_list, zuv_list = [], []
        for table in rpd_doc.tables:
            if len(table.columns) != 5:
                continue
            words = 'Планируемые результаты обучения по дисциплине'.split()
            if all(map(lambda word: word in table.cell(0, 3).text, words)):
                for cell in list(table.column_cells(1))[1:]:
                    comp_list.append(cell.text)
                for cell in list(table.column_cells(3))[1:]:
                    zuv_list.append(cell.text)
        if len(comp_list) == len(zuv_list):
            competences = {}
            for comp, zuv_str in zip(comp_list, zuv_list):
                comp_words = comp.split()
                if comp_words:
                    comp_code = comp_words[0]
                    if comp_code.endswith('.') or comp_code.endswith(':'):
                        comp_code = comp_code[:-1]
                    competences[comp_code] = core.get_zuv(zuv_str)
            result[code] = core.RPD(code=code, name=name, zuv=core.ZUV(), competences=competences)
        # elif len(zuv_list) == 1:
        #     print(filename, core.get_zuv(zuv_list[0]))
        #     zuv = core.ZUV()
        #     result[code] = core.RPD(code=code, name=name, )
        else:
            print(filename)
    return result


def main(args: Namespace) -> None:
    """ Точка входа """
    plan = core.get_plan(args.plan)
    rpd_dict = get_rpd_dict(plan, args.rpd_dir)
    template = core.get_template('fos.docx')
    context = {
        'plan': plan,
        'rpd_dict': rpd_dict
    }
    fill_table_1(template, context)
    # fill_table_2_1(template, context)
    # fill_table_4(template, context)
    template.render(context)
    template.save(args.plan[:-4] + '.docx')
    print('Partially done')


if __name__ == '__main__':
    parser = ArgumentParser()
    parser.add_argument('plan', help='PLX-файл РУПа')
    parser.add_argument('rpd_dir', help='каталог РПД')
    main(parser.parse_args())
