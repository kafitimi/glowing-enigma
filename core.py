""" Базовые классы """
import sys
from os.path import join
from typing import Dict, List, NamedTuple

from docx.table import Table
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxtpl import DocxTemplate

BACHELOR = 2
MASTER = 3

CENTER = 'Table Heading'
JUSTIFY = 'Table Contents'


def get_template(filename: str) -> DocxTemplate:
    """ Читаем шаблон РПД """
    try:
        template = DocxTemplate(join('templates', filename))
    except OSError:
        print('Не могу открыть шаблон')
        sys.exit()
    return template


def set_cell_text(table: Table, row: int, col: int, style: str, text: str) -> None:
    """ Добавить текст в ячейку таблицы """
    cell = table.cell(row, col)
    if cell.text:
        cell.add_paragraph(text, style)
    else:
        cell.text = text
        cell.paragraphs[0].style = style


def add_table_rows(table: Table, rows: int) -> None:
    """ Добавить строки в таблицу """
    def get_border(name):
        border = OxmlElement(f'w:{name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        return border

    def get_borders():
        borders = OxmlElement("w:tcBorders")
        borders.append(get_border('top'))
        borders.append(get_border('left'))
        borders.append(get_border('bottom'))
        borders.append(get_border('right'))
        return borders

    for _ in range(rows):
        row = table.add_row()
        for cell in row.cells:
            props = cell._element.tcPr  # pylint: disable=protected-access
            props.append(get_borders())


def remove_table(template: DocxTemplate, table_index: int) -> None:
    """ Удаляем таблицу из шаблона по её индексу """
    docx = template.get_docx()
    table_element = docx.tables[table_index]._element  # pylint: disable=protected-access
    parent_element = table_element.getparent()
    parent_element.remove(table_element)


class ZUV(NamedTuple):
    raw: str = ''
    knowledge: List[str] = []
    abilities: List[str] = []
    skills: List[str] = []


def get_zuv(content: str) -> ZUV:
    rules1 = [
        ('знать:', 'knowledge'), ('знает:', 'knowledge'), ('знать', 'knowledge'), ('знает', 'knowledge'),
        ('уметь:', 'abilities'), ('умеет:', 'abilities'), ('уметь', 'abilities'), ('умеет', 'abilities'),
        ('владеть:', 'skills'), ('владеет:', 'skills'), ('владеть', 'skills'), ('владеет', 'skills'),
    ]
    rules2 = [
        ('методиками', 'методиками '), ('методиками:', 'методиками '),
        ('навыками:', 'навыками '), ('навыками', 'навыками '),
        ('практическими навыками', 'практическими навыками '), ('практическими навыками:', 'практическими навыками '),
        ('опытом', 'опытом '), ('опытом:', 'опытом '),
        ('практическим опытом', 'практическим опытом '), ('практическим опытом:', 'практическим опытом '),
    ]
    zuv = {'raw': content, 'knowledge': [], 'abilities': [], 'skills': []}
    current, prefix = '', ''
    for line in content.splitlines():
        line = line.strip()
        simple = line.lower()
        for trigger, action in rules1:
            if simple.startswith(trigger):
                line = line[len(trigger):].strip()
                current, prefix = action, ''
                break
        if current == 'skills':
            simple = line.lower()
            for trigger, action in rules2:
                if simple.startswith(trigger):
                    line = line[len(trigger):].strip()
                    prefix = action
                    break
        if current and line:
            zuv[current].append(prefix + line)
    return ZUV(**zuv)


class RPD(NamedTuple):
    code: str = ''
    name: str = ''
    zuv: ZUV = ZUV()
    competences: Dict[str, ZUV] = {}
