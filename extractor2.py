""" Генерация ФОС """
import difflib
import os
import sys
from typing import Any, Dict, List

import numpy as np
import pandas as pd
from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import CT_P, CT_Tbl
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docxtpl import DocxTemplate

import core
from enigma import Competence, EducationPlan, Subject, get_plan, word_doc
from enigma.education_plan import CT_EXAM, CT_COURSEWORK, CT_CREDIT, CT_CREDIT_GRADE
from enigma.word_doc import add_table_rows, set_cell_text


def iterate_items(parent: Any):
    """ Обход параграфов и таблиц в документе """
    if isinstance(parent, DocumentType):
        parent_elem = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elem = parent._tc
    else:
        raise ValueError('Oops')

    for child in parent_elem.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def get_section_paragraphs(input_filename: str, start_kw: List[str], final_kw: List[str]) -> List[str]:
    """ Извлечь список абзацев текста из docx-файла """
    result, source = Document(), Document(input_filename)
    started = False
    paragraphs = []
    for item in iterate_items(source):
        if not started:
            if isinstance(item, Paragraph) and any(kw in item.text for kw in start_kw):
                started = True
        else:
            if isinstance(item, Paragraph) and any(kw in item.text for kw in final_kw):
                break
            if isinstance(item, Paragraph):
                text = item.text.strip()
                if text:
                    paragraphs.append(text + '\n')
            elif isinstance(item, Table):
                pass
    return paragraphs


def get_rpd(name):
    result = None
    for fn in os.listdir('rpds'):
        if fn.endswith('.docx'):
            if name in fn:
                result = os.path.join('rpds', fn)
                break
    return result


def check_args() -> None:
    """ Проверка аргументов командной строки """
    if len(sys.argv) != 3:
        print('Синтаксис:\n\tpython {0} <руп> <фос>'.format(*sys.argv))
        sys.exit()


def fill_table_1(template: DocxTemplate, context: Dict[str, any]) -> None:
    """ Заполнение таблиц с формами контроля """
    control_fancy_name = {
        CT_EXAM: 'Экзамен',
        CT_CREDIT_GRADE: 'Зачет с оценкой',
        CT_CREDIT: 'Зачет',
        CT_COURSEWORK: 'Курсовой проект',
    }

    plan: EducationPlan = context['plan']
    if plan.degree == core.BACHELOR:
        word_doc.remove_table(template, 1)
    elif plan.degree == core.MASTER:
        word_doc.remove_table(template, 2)
    table: Table = template.get_docx().tables[1]

    row_number = 0
    for competence in sorted(plan.competence_codes.values(), key=Competence.repr):
        add_table_rows(table, 1)
        row = len(table.rows) - 1
        row_number += 1
        set_cell_text(table, row, 0, word_doc.CENTER, str(row_number))
        set_cell_text(table, row, 1, word_doc.JUSTIFY, competence.code + ' ' + competence.description)
        table.cell(row, 1).merge(table.cell(row, len(table.columns) - 1))
        subjects = [plan.subject_codes[s] for s in competence.subjects]
        for subject in sorted(subjects, key=Subject.repr):
            add_table_rows(table, 1)
            row = len(table.rows) - 1
            row_number += 1
            set_cell_text(table, row, 0, word_doc.CENTER, str(row_number))
            set_cell_text(table, row, 1, word_doc.JUSTIFY, subject.code + ' ' + subject.name)
            for number, semester in subject.semesters.items():
                controls = [control_fancy_name[c] for c in semester.control]
                set_cell_text(table, row, number+1, word_doc.CENTER, ', '.join(controls))


def table2list(x):
    """ Читает таблицу из РПД в список """
    ls = []
    try:
        for row in x.rows:
            for cell in row.cells:
                celltext = ''
                for paragraph in cell.paragraphs:
                    celltext += '\n' + paragraph.text
                ls.append(celltext)
        return [ls, len(x.rows), len(x.columns)]
    except IndexError:
        return [[' '], 1, 1]


cases = {
    'отл': 'Отлично',
    'хор': 'Хорошо',
    'неудов': 'неУдов',
    'удовл': 'Удовл',
    'не зач': 'не Зачтено',
    'незач': 'не Зачтено',
    'зач': 'Зачтено'
}

testtypes = {
    'the': 'Перевод, лексико-грамматический анализ, объем словарного запаса и аннотирование научной статьи',
    'экзамен': 'Экзаменационные билеты',
    'зачет': 'Задания к зачету'
}
       

def normalize(alist):
    thelist = []
    for item in alist:
        item = item.lower()
        for key, value in cases.items():
            if key in item:
                item = value
        if item not in cases.values():
            item = ' '
        thelist.append(item)
    return thelist


def findout_fos_part(documento):
    """ Выдергивает все абзацы между литературой и "Примерные контрольные задания" """
    text_at_start = ['Примерные контрольные задания (вопросы']
    text_at_end = ['учебной литературы', 'iprbook', 'lanbook', 'НБ СВФУ',
                   'информационно-телекоммуникационной сети']
    flag, result_text = False, ''
    for paragraph in documento.paragraphs:
        if any(key_word in paragraph.text for key_word in text_at_end):
            flag = False
        if flag:
            result_text += (paragraph.text+'\n')
        if text_at_start[0] in paragraph.text:
            flag = True
    result_text = result_text.replace('\n\n', '\n').replace('  ', ' ', 1000)
    return result_text


def find_rpd(subjectcode, subjectname, control, controls, sem):
    global fileslist
    filename = difflib.get_close_matches(subjectcode + ' ' + subjectname, fileslist)
    if len(filename) < 1:
        return [' ', ' ', ' ', ' ', ' ', ' ']
    filename = './rpds/' + str(filename[0]) + '.docx'
    rpd_doc = Document(filename)
    marks = []
    crirs = []
    zuv = ' '   
    zuv_not_found, zuv_not_found2 = True, True
    for table in rpd_doc.tables:
        tablen = table2list(table)
        row, column = tablen[1], tablen[2]
        tablen = tablen[0]
        if len(tablen) < column:
            continue
        else: 
            row = len(tablen) // column
            tablen = tablen[:(row*column)]
        df = pd.DataFrame(np.array(tablen).reshape(row, column))  # reshape to the table shape

        if sum([key_word in df.iloc[0, column-1] for key_word in ['Шкал', 'Оценк']]) > 0:
            marks += list(df[df.columns[-1]])
            crirs += list(df[df.columns[-2]])
        try:
            if zuv_not_found:
                for irow in range(3, 0, -1):
                    for icol in range(1, 3):
                        if '(ЗУВ)' in df.iloc[irow, icol]:
                            zuv += df.iloc[irow + 1, icol]
                            zuv_not_found = False

            if zuv_not_found2 and zuv_not_found:
                for irow in range(3, 0, -1):
                    for icol in range(1, 3):
                        if '.1.2.' in df.iloc[irow, icol]:
                            zuv += df.iloc[irow+1, icol]
                            zuv_not_found2 = False
        except IndexError: 
            cndekc = 7

    marks = normalize(marks)
    crirs = [x for ind, x in enumerate(crirs) if marks[ind] != ' ']
    marks = [x for x in marks if x != ' ']
    excellent, good, fair, bad, tests = ' ', ' ', ' ', ' ', ' ' 
    if control == 'Зачет':
        for ind, criteria in enumerate(crirs):
            if marks[ind] == 'Зачтено': excellent += crirs[ind]
            if marks[ind] == 'не Зачтено': bad += crirs[ind]
    else:
        for ind, criteria in enumerate(crirs):
            if marks[ind] == 'Отлично': excellent += crirs[ind]
            if marks[ind] == 'Хорошо': good += crirs[ind]
            if marks[ind] == 'Удовл': fair += crirs[ind]
            if marks[ind] == 'не Удов': bad += crirs[ind]
    
    if bad == ' ': bad = excellent
    document = Document(filename)
    bigtext = findout_fos_part(document) 
    bigtext = (control + bigtext).lower()
    testtype = ' '
    for key, value in testtypes.items():
        if key in bigtext:
            testtype = value
            continue

    return [zuv.replace('\n\n', '\n'), excellent, good, fair, bad, testtype]
 

def fill_table_2(template: DocxTemplate, context: Dict[str, any]) -> None:
    """ Заполнение таблиц с формами контроля """
    global fileslist 
    control_fancy_name = {
        CT_EXAM: 'Экзамен',
        CT_CREDIT_GRADE: 'Зачет с оценкой',
        CT_CREDIT: 'Зачет',
        CT_COURSEWORK: 'Курсовой проект',
    }

    plan: EducationPlan = context['plan']
    table: Table = template.get_docx().tables[2]
    fileslist = [filename[:-5] for filename in os.listdir('./rpds') if filename.endswith('.docx')]

    row_number = 0
    for competence in sorted(plan.competence_codes.values(), key=Competence.repr):
        add_table_rows(table, 1)
        row = len(table.rows) - 1
        row_number += 1
        set_cell_text(table, row, 0, word_doc.CENTER, str(row_number))
        set_cell_text(table, row, 1, word_doc.JUSTIFY, competence.code + ' ' + competence.description)
        for runover in range(2, 8):
            set_cell_text(table, row, runover, word_doc.JUSTIFY, ' ')
        # table.cell(row, 1).merge(table.cell(row, len(table.columns) - 1))
        subjects = [plan.subject_codes[s] for s in competence.subjects]
        for subject in sorted(subjects, key=Subject.repr):
            # add_table_rows(table, 1)
            # row = len(table.rows) - 1
            # len(table.rows) - 1
            # row_number += 1
            controls = []
            for number, semester in subject.semesters.items():
                controls += [control_fancy_name[c] for c in semester.control]
             
            sem = 0
            for control in controls: 
                add_table_rows(table, 1)
                row = len(table.rows) - 1             
                zuv_criteria = find_rpd(subject.code, subject.name, control, controls, sem)
                sem += 1
                if len(zuv_criteria) == 6:
                    set_cell_text(table, row, 0, word_doc.CENTER, ' ')
                    set_cell_text(table, row, 1, word_doc.JUSTIFY, subject.code + ' ' + subject.name)
                    set_cell_text(table, row, 2, word_doc.JUSTIFY, zuv_criteria[0])
                    if control == 'Зачет':
                        set_cell_text(table, row, 3, word_doc.JUSTIFY,  zuv_criteria[1])
                        set_cell_text(table, row, 4, word_doc.JUSTIFY, ' ')
                        set_cell_text(table, row, 5, word_doc.JUSTIFY, ' ')
                        set_cell_text(table, row, 6, word_doc.JUSTIFY, zuv_criteria[4])
                        table.cell(row, 3).merge(table.cell(row, 5))                
                    else:
                        set_cell_text(table, row, 3, word_doc.JUSTIFY, zuv_criteria[1])
                        set_cell_text(table, row, 4, word_doc.JUSTIFY, zuv_criteria[2])
                        set_cell_text(table, row, 5, word_doc.JUSTIFY, zuv_criteria[3])
                        set_cell_text(table, row, 6, word_doc.JUSTIFY, zuv_criteria[4])
                    set_cell_text(table, row, 7, word_doc.JUSTIFY, zuv_criteria[5])


def fill_table_2_1(template: DocxTemplate, context: Dict[str, any]) -> None:
    """ Заполнение таблицы в разделе 2.1 """
    plan: EducationPlan = context['plan']
    table: Table = template.get_docx().tables[3]
    for subject in sorted(plan.subject_codes.values(), key=Subject.repr):
        add_table_rows(table, 1)
        row_index = len(table.rows) - 1
        set_cell_text(table, row_index, 0, word_doc.CENTER, subject.code)
        set_cell_text(table, row_index, 1, word_doc.JUSTIFY, subject.name)


header = [
    'Министерство науки и высшего образования Российской Федерации',
    'Федеральное государственное автономное образовательное учреждение',
    'высшего образования',
    '«СЕВЕРО-ВОСТОЧНЫЙ ФЕДЕРАЛЬНЫЙ УНИВЕРСИТЕТ ИМЕНИ М.К. АММОСОВА»',
    'Институт математики и информатики',
    '\n\n\n\n\n\n\n\n\n\n\n\nФонд оценочных средств дисциплины'
]

middle = [
    '\nкод дисциплины и название дисциплины',  # bold
    '\nдля программы ',
    'по направлению подготовки',
    'УГС Название направления подготовки — ',
    'Направленность программы: ',
    '\n\nФорма контроля:',  # bold
    'экзамен/зачет/зачет со оценкой/курсовая работа',  # bold
    '\n\n'
]

middleBold = [True, False, False, False, False, True, True, False]

footer = ['\n\n\n\n\n\n\n\n\n\n\n\n\nЯкутск 2019']

method_mater = [
    '6.3. Методические материалы, определяющие процедуры оценивания',
    'Текущий контроль успеваемости с БРС состоит из:',
    '– контрольные срезы, устанавливаемые учебными подразделениями;',
    '– рубежный срез, проводимый после окончания зачетной недели, за день до начала',
    'экзаменационной сессии, который определяет общую сумму баллов, набранную '
    'обучающимся в течение всего семестра по данной дисциплине.',
    'Промежуточная аттестация осуществляется в конце семестра после завершения учебных '
    'занятий в форме: зачетов, зачетов с оценкой, оценки отчета по практике, экзаменов.',
    'Оценка знаний, умений и навыков производится согласно положению о балльно-рейтинговой системе в СВФУ.'
]

table61 = [
    'Коды', 'оцениваемых', 'Индикаторы', 'достижения',
    'Показатель', 'оценивания', '(по', 'п.1.2.РПД)', 'Шкалы оценивания',
    'сформированности'
]

table62 = [
    'Оцениваемый', 'показатель', '(ЗУВ)', 'Тема', '(темы)',
    'Образец', 'типового', 'тестового', 'задания', '(вопроса)'
]

tableBooks = [
    'Автор', 'название', 'издания', 'издательство', 'литературы',
    'информационных ресурсов', 'гриф', 'НБ СВФУ', 'библиотека',
    'экземпляров', 'Электронные', 'ЭБС', 'ЭБ СВФУ'
]


def preceding_paragraph(fos_doc, x):
    """ Определяет тип таблицы и добавляет название параграфа """
    global table_flazhok 
    if sum([word in x for word in table61]) > 5:
        title6 = '6. Фонд оценочных средств для проведения промежуточной аттестации обучающихся по дисциплине'
        fos_doc.add_paragraph(title6).bold = True
        fos_doc.paragraphs[-1].runs[0].bold = True 
        fos_doc.add_paragraph('6.1. Показатели, критерии и шкала оценивания').bold = True
        fos_doc.paragraphs[-1].runs[0].bold = True 
        table_flazhok = True
        return True
    if sum([word in x for word in table62]) > 5:
        fos_doc.add_paragraph('6.2. Примерные контрольные задания (вопросы) для промежуточной аттестации').bold = True
        fos_doc.paragraphs[-1].runs[0].bold = True 
        table_flazhok = True
        return True
    if sum([word in x for word in tableBooks]) > 5:
        table_flazhok = False
    return table_flazhok


def findout_fos(documento):
    """ Выдергивает все абзацы между литературой и "Примерные контрольные задания" """
    text_at_start = ['Примерные контрольные задания (вопросы']
    text_at_end = ['учебной литературы', 'iprbook', 'lanbook', 'НБ СВФУ', 'информационно-телекоммуникационной сети']
    flag, result_text = False, ''
    for paragraph in documento.paragraphs:
        if any(key_word in paragraph.text for key_word in text_at_end):
            flag = False
        if flag:
            result_text += (paragraph.text + '\n')
        if text_at_start[0] in paragraph.text:
            flag = True
    result_text = result_text.replace('\n\n', '\n').replace('  ', ' ', 1000)
    return result_text


def list2docx(fos_doc, ls, row, column):
    """ Из списка восстанавливает таблицу, объединяя ячейки со совпадающими текстами """
    if len(ls) != row * column:
        return 

    df = pd.DataFrame(np.array(ls).reshape(row, column))  # reshape to the table shape
    word_table = fos_doc.add_table(rows=row, cols=column, style='Table Grid')
    for ind_row in range(0, row, 1):
        for ind_col in range(0, column, 1):
            cell = word_table.cell(ind_row, ind_col)
            if ind_row > 0:
                cell2 = word_table.cell(ind_row-1, ind_col)
                if cell2.text == df.iloc[ind_row, ind_col]:
                    cell.text = ''
                    cell.merge(cell2)
                else:
                    cell.text = df.iloc[ind_row, ind_col]
            else:
                if ind_col > 0:
                    cell2 = word_table.cell(ind_row, ind_col-1)
                    if cell2.text == df.iloc[ind_row, ind_col]:
                        cell.text = ''
                        cell.merge(cell2)
                    else:
                        cell.text = df.iloc[ind_row, ind_col]
    fos_doc.add_paragraph(' ')                                      
    return 


def fill_section_2_2(template: DocxTemplate, context: Dict[str, any]) -> None:
    global table_flazhok
    control_fancy_name = {
        CT_EXAM: 'Экзамен',
        CT_CREDIT_GRADE: 'Зачет с оценкой',
        CT_CREDIT: 'Зачет',
        CT_COURSEWORK: 'Курсовой проект',
    }

    """ Заполнение раздела 2.2 """
    # marker = None
    # for p1 in template.get_docx().paragraphs:
    #     keywords = ['оценочные средства для', 'государственной итоговой аттестации']
    #     if all(kw in p1.text.lower() for kw in keywords):
    #         marker = p1
    #         break
    global fileslist
    fileslist = [filename[:-5] for filename in os.listdir('./rpds') if filename[-4:] == 'docx']

    plan: EducationPlan = context['plan']
    middle[1] += 'магистратуры' if plan.degree == core.MASTER else 'бакалавриата'
    middle[4] += plan.program
    middle[3] = plan.code + ' ' + plan.name
    subjects = sorted(plan.subject_codes.values(), key=Subject.repr)
    for s in subjects:
        rpd = difflib.get_close_matches(s.code + ' ' + s.name, fileslist)
        if len(rpd) < 1:
            continue
        rpd = './rpds/' + str(rpd[0]) + '.docx'
        rpd_doc = Document(rpd)

        # титульная страница
        document = template.get_docx()
        document.add_page_break() 
        document.add_paragraph('\n'.join(header))
        document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        controls = []
        for number, semester in s.semesters.items():
            controls += [control_fancy_name[c] for c in semester.control]
        middle[0] = s.code + ' ' + s.name
        middle[6] = ' / '.join(controls)
        for ind, line in enumerate(middle):
            document.add_paragraph(line)
            document.paragraphs[-1].runs[0].bold = middleBold[ind]
            document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_paragraph('\n'.join(footer)) 
        document.paragraphs[-1].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        document.add_page_break() 

        ''' флажок, чтобы забрать все таблицы между двумя группами ключевых слов '''
        table_flazhok = False
        for table in rpd_doc.tables:
            tablen = table2list(table)
            if preceding_paragraph(document, ' '.join(tablen[0][:6])):
                list2docx(document, *tablen)
        ''' таблицы все включены, теперь забираем все абзацы '''
        text_heap = findout_fos(rpd_doc)
        document.add_paragraph(text_heap) 
        if 'Методические материалы, определяющие' not in text_heap:
            document.add_paragraph('\n'.join(method_mater))

        # p2 = marker.insert_paragraph_before('%s %s' % (s.code, s.name))
        # p2.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # p2.paragraph_format.first_line_indent = Cm(0)
        # for r in p2.runs:
        #     r.bold = True

        # for item in iterate_items(Document(rpd)):
        #     if isinstance(item, Paragraph):
        #         p3 = marker.insert_paragraph_before(item.text)
        #         p3.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        #         p3.paragraph_format.first_line_indent = Cm(0)

    # plan: core.EducationPlan = context['plan']
    # table: Table = document.get_docx().tables[3]
    # for subject in sorted(plan.subject_codes.values(), key=core.Subject.repr):
    #     add_table_rows(table, 1)
    #     row_index = len(table.rows) - 1
    #     set_cell_text(table, row_index, 0, word_doc.CENTER, subject.code)
    #     set_cell_text(table, row_index, 1, word_doc.JUSTIFY, subject.name)


def fill_table_4(template: DocxTemplate, context: Dict[str, any]) -> None:
    """ Заполнение бланка "Лист сформированности компетенций" """
    plan: EducationPlan = context['plan']
    table: Table = template.get_docx().tables[-1]
    row_number = 0
    for competence in sorted(plan.competence_codes.values(), key=Competence.repr):
        add_table_rows(table, 1)
        row_index = len(table.rows) - 1
        row_number += 1
        set_cell_text(table, row_index, 0, word_doc.CENTER, str(row_number))
        set_cell_text(table, row_index, 1, word_doc.JUSTIFY, competence.code + ' ' + competence.description)
        subjects = [plan.subject_codes[s] for s in competence.subjects]
        for subject in sorted(subjects, key=Subject.repr):
            add_table_rows(table, 1)
            row_index = len(table.rows) - 1
            set_cell_text(table, row_index, 1, word_doc.JUSTIFY, subject.code + ' ' + subject.name)

    add_table_rows(table, 1)
    row_number += 1
    row_index = len(table.rows) - 1
    set_cell_text(table, row_index, 0, word_doc.CENTER, str(row_number))
    set_cell_text(table, row_index, 1, word_doc.JUSTIFY, 'Практики')

    add_table_rows(table, 1)
    row_number += 1
    row_index = len(table.rows) - 1
    set_cell_text(table, row_index, 0, word_doc.CENTER, str(row_number))
    set_cell_text(table, row_index, 1, word_doc.JUSTIFY, 'НИР')


def main() -> None:
    """ Точка входа """
    check_args()
    global fileslist

    plan = get_plan(sys.argv[1])
    template = word_doc.get_template('fos.docx')
    context = {
        'plan': plan,
    }
    fill_table_1(template, context)
    fileslist = ' '
    fill_table_2(template, context)
    fill_table_2_1(template, context)
    fill_table_4(template, context)
    fill_section_2_2(template, context)
    template.render(context)
    template.save('filled_'+sys.argv[2])
    print('Partially done')


if __name__ == '__main__':
    table_flazhok = False
    fileslist = ' '
    main()
