""" Генерация РПД """
import argparse
import functools
import glob
import os
import sys
from operator import itemgetter
from typing import List, Dict, Any

from Levenshtein import distance as levenshtein_d  # pylint: disable=no-name-in-module
from docx.shared import Mm
from docx.table import Table, _Row
from docxtpl import DocxTemplate, InlineImage

from enigma import Course, Competence, EducationPlan, Subject, get_plan, word_doc
from enigma.education_plan import CT_EXAM, CT_CREDIT, CT_CREDIT_GRADE
from enigma.word_doc import add_table_rows, set_cell_text

IMAGE_KINDS = ('lit', 'title')


def fill_table_column(table: Table, row: int, columns: List[int], values: List[Any]) -> None:
    """ Заполнить колонку таблицу """
    for value in values:
        str_value = str_or_dash(value)
        for col in columns:
            style = word_doc.CENTER if len(str_value) < 5 else word_doc.JUSTIFY
            set_cell_text(table, row, col, style, str_value)
        row += 1


def distribute(total: int, portions: int) -> List[int]:
    """ Распределить часы по темам """
    base, remainder = divmod(total, portions)
    return [base] * (portions - remainder) + [base + 1] * remainder


def str_or_dash(value: Any) -> str:
    """ Конвертируем целое в строку """
    return str(value) if value else '—'


def check_args() -> None:
    """ Проверка аргументов командной строки """
    if len(sys.argv) != 3:
        print('Синтаксис:\n\tpython {0} <руп> <курс>'.format(*sys.argv))
        sys.exit()


def get_course(course_filename: str) -> Course:
    """ Открываем курс обучения """
    try:
        course = Course(course_filename)
    except OSError:
        print('Не могу открыть курс обучения' % course_filename)
        sys.exit()
    return course


def get_subject(plan: EducationPlan, course: Course) -> Subject:
    """ Ищем подходящую дисциплину в учебном плане """
    result = plan.find_subject(course.names)
    if result is None:
        print('Не могу найти подходящую дисциплину в учебном плане')
        sys.exit()
    return result


def fill_table_1_2(template: DocxTemplate, context: Dict[str, Any]) -> None:
    """ Заполняем таблицу с компетенциями в разделе 1.2 """
    table = template.get_docx().tables[1]
    add_table_rows(table, len(context['subject'].competencies))

    row = 0
    competencies = [context['plan'].competence_codes[c] for c in context['subject'].competencies]
    for competence in sorted(competencies, key=Competence.repr):
        row += 1
        set_cell_text(table, row, 0, word_doc.CENTER, competence.category)
        set_cell_text(table, row, 1, word_doc.JUSTIFY, competence.code + ' ' + competence.description)
        set_cell_text(table, row, 4, word_doc.CENTER, context['course'].assessment)

        for ind_code in sorted(competence.indicator_codes):
            indicator = competence.indicator_codes[ind_code]
            set_cell_text(table, row, 2, word_doc.JUSTIFY, ind_code + ' ' + indicator.description)

    def add_study_results(attr: str, caption: str) -> None:
        """ Добавить в ячейку таблицы результаты обучения """
        results = context['course'].__getattribute__(attr)
        if results:
            set_cell_text(table, 1, 3, word_doc.JUSTIFY, caption)
            for elem in results:
                set_cell_text(table, 1, 3, 'Table List', '•\t' + elem)

    table.cell(1, 3).merge(table.cell(row, 3))
    add_study_results('knowledge', 'Знать:')
    add_study_results('abilities', 'Уметь:')
    add_study_results('skills', 'Владеть:')


def fill_table_2(template: DocxTemplate, context: Dict[str, Any]) -> None:
    """ Убираем из таблицы с выпиской из учебного плана в разделе 2 лишние строки """
    def remove_row(tbl: Table, row_index: int):
        row: _Row = tbl.rows[row_index]
        tbl._tbl.remove(row._tr)

    table = template.get_docx().tables[3]
    subject: Subject = context['subject']
    practical_hours = subject.get_practical_hours()
    if not practical_hours:
        remove_row(table, 18)
        remove_row(table, 15)
        remove_row(table, 13)
        remove_row(table, 10)


def fill_table_3_1(template: DocxTemplate, context: Dict[str, Any]) -> None:
    """ Заполняем таблицу с содержанием курса в разделе 3.1 """

    # Извлекаем названия тем и форматируем для заполнения таблицы
    course, subject = context['course'], context['subject']
    themes = ['Тема %d. %s' % (i + 1, t['тема']) for i, t in enumerate(course.themes)]

    # Равномерно размазываем часы по темам
    themes_count = len(themes)
    lectures = distribute(subject.get_hours('lectures'), themes_count)
    practices = distribute(subject.get_hours('practices'), themes_count)
    labworks = distribute(subject.get_hours('labworks'), themes_count)
    controls = distribute(subject.get_hours('controls'), themes_count)
    homeworks = distribute(subject.get_hours('homeworks'), themes_count)
    totals = [sum(t) for t in zip(lectures, practices, labworks, controls, homeworks)]

    # Заполняем таблицу
    table = template.get_docx().tables[4]
    add_table_rows(table, themes_count + 1)  # последняя строка - для итога
    fill_table_column(table, 2, [0], themes + ['Всего часов'])
    fill_table_column(table, 2, [1], totals + [subject.get_total_hours()])
    fill_table_column(table, 2, [2], lectures + [sum(lectures)])
    fill_table_column(table, 2, [4], labworks + [sum(labworks)])
    fill_table_column(table, 2, [6], practices + [sum(practices)])
    fill_table_column(table, 2, [8], controls + [sum(controls)])
    fill_table_column(table, 2, [9], homeworks + [sum(homeworks)])
    fill_table_column(table, 2, [3, 5, 7], [0] * (themes_count + 1))  # пустые значения


def fill_table_4(template: DocxTemplate, context: Dict[str, Any]) -> None:
    """ Заполняем таблицу с содержанием СРС в разделе 4 """
    themes = context['course'].themes
    themes_count = len(themes)
    subject = context['subject']
    homeworks = distribute(subject.get_hours('homeworks'), themes_count)

    table = template.get_docx().tables[5]
    add_table_rows(table, themes_count + 1)  # последняя строка - для итога

    hw_text = "Проработка теоретического материала, подготовка к выполнению лабораторной работы"
    hw_control = "Вопросы к итоговому тесту"

    i = 0
    for theme in themes:
        i += 1
        set_cell_text(table, i, 0, word_doc.CENTER, str(i))
        set_cell_text(table, i, 1, word_doc.CENTER, theme['тема'])
        set_cell_text(table, i, 2, word_doc.CENTER, hw_text)
        set_cell_text(table, i, 3, word_doc.CENTER, str_or_dash(homeworks[i - 1]))
        set_cell_text(table, i, 4, word_doc.CENTER, hw_control)

    i += 1
    set_cell_text(table, i, 1, word_doc.JUSTIFY, 'Всего часов')
    set_cell_text(table, i, 3, word_doc.CENTER, str_or_dash(sum(homeworks)))


def fill_table_6_1(template: DocxTemplate, context: Dict[str, Any]):
    """ Заполняем таблицу в разделе 6.1 """

    def add_study_results(attr: str, caption: str, row: int, col: int) -> None:
        """ Добавить в ячейку таблицы результаты обучения """
        results = course.__getattribute__(attr)
        if results:
            set_cell_text(table, row, col, word_doc.JUSTIFY, caption)
            for elem in results:
                set_cell_text(table, row, col, 'Table List', '•\t' + elem)

    course, subject = context['course'], context['subject']

    # Уровни освоения
    control = [s.control for s in subject.semesters.values()]
    control = functools.reduce(lambda x, y: x + y if isinstance(x, list) else set.union(x, y), control)
    if CT_EXAM in control:
        levels = [
            ('Высокий', 'Отлично'), ('Базовый', 'Хорошо'),
            ('Минимальный', 'Удовлетворительно'), ('Не освоены', 'Неудовлетворительно'),
        ]
    elif CT_CREDIT_GRADE in control:
        levels = [
            ('Высокий', 'Зачтено (отлично)'), ('Базовый', 'Не зачтено (хорошо)'),
            ('Минимальный', 'Зачтено (удовлетворительно)'), ('Не освоены', 'Не зачтено'),
        ]
    else:
        levels = [('Освоено', 'Зачтено'), ('Не освоено', 'Не зачтено')]

    # Строки таблицы
    table = template.get_docx().tables[8]
    rows_count = len(subject.competencies) * len(levels)
    add_table_rows(table, rows_count)

    # Компетенции и индикаторы
    start_row = 2
    for code in subject.competencies:
        competence = context['plan'].competence_codes[code]
        table.cell(start_row, 0).merge(table.cell(start_row + len(levels) - 1, 0))
        set_cell_text(table, start_row, 0, word_doc.JUSTIFY, competence.code + ' ' + competence.description)
        table.cell(start_row, 1).merge(table.cell(start_row + len(levels) - 1, 1))
        for ind_code in sorted(competence.indicator_codes):
            indicator = competence.indicator_codes[ind_code]
            set_cell_text(table, start_row, 1, word_doc.JUSTIFY, ind_code + ' ' + indicator.description)
        start_row += len(levels)

    # Знать, уметь, владеть
    start_row = 2
    table.cell(start_row, 2).merge(table.cell(start_row + rows_count - 1, 2))
    add_study_results('knowledge', 'Знать:', 2, 2)
    add_study_results('abilities', 'Уметь:', 2, 2)
    add_study_results('skills', 'Владеть:', 2, 2)

    # Уровни освоения
    start_row = 2
    for level, grade in levels:
        table.cell(start_row, 3).merge(table.cell(start_row + len(subject.competencies) - 1, 3))
        set_cell_text(table, start_row, 3, word_doc.CENTER, level)
        table.cell(start_row, 4).merge(table.cell(start_row + len(subject.competencies) - 1, 4))
        if CT_CREDIT in control:
            if level == 'Освоено':
                add_study_results('knowledge', 'Обучаемый знает:', start_row, 4)
                add_study_results('abilities', 'Обучаемый умеет:', start_row, 4)
                add_study_results('skills', 'Обучаемый владеет:', start_row, 4)
            else:
                add_study_results('knowledge', 'Обучаемый не знает:', start_row, 4)
                add_study_results('abilities', 'Обучаемый не умеет:', start_row, 4)
                add_study_results('skills', 'Обучаемый не владеет:', start_row, 4)
        else:
            if level == 'Высокий':
                add_study_results('knowledge', 'Обучаемый знает:', start_row, 4)
                add_study_results('abilities', 'Обучаемый умеет:', start_row, 4)
                add_study_results('skills', 'Обучаемый владеет:', start_row, 4)
            elif level == 'Базовый':
                add_study_results('knowledge', 'Обучаемый знает:', start_row, 4)
                add_study_results('abilities', 'Обучаемый не умеет:', start_row, 4)
                add_study_results('skills', 'Обучаемый владеет:', start_row, 4)
            elif level == 'Минимальный':
                add_study_results('knowledge', 'Обучаемый не знает:', start_row, 4)
                add_study_results('abilities', 'Обучаемый не умеет:', start_row, 4)
                add_study_results('skills', 'Обучаемый владеет:', start_row, 4)
            else:
                add_study_results('knowledge', 'Обучаемый не знает:', start_row, 4)
                add_study_results('abilities', 'Обучаемый не умеет:', start_row, 4)
                add_study_results('skills', 'Обучаемый не владеет:', start_row, 4)
        table.cell(start_row, 5).merge(table.cell(start_row + len(subject.competencies) - 1, 5))
        set_cell_text(table, start_row, 5, word_doc.CENTER, grade)
        start_row += len(subject.competencies)


def fill_table_6_2(template: DocxTemplate, context: Dict[str, Any]) -> None:
    def add_study_results(attr: str, caption: str) -> None:
        """ Добавить в ячейку таблицы результаты обучения """
        results = context['course'].__getattribute__(attr)
        if results:
            set_cell_text(table, 1, 2, word_doc.JUSTIFY, caption)
            for elem in results:
                set_cell_text(table, 1, 2, 'Table List', '•\t' + elem)

    table = template.get_docx().tables[9]
    course = context['course']
    subject = context['subject']
    add_table_rows(table, len(subject.competencies))

    # Компетенции и индикаторы
    row = 0
    competencies = [context['plan'].competence_codes[c] for c in context['subject'].competencies]
    for competence in sorted(competencies, key=Competence.repr):
        row += 1
        set_cell_text(table, row, 0, word_doc.JUSTIFY, competence.code + ' ' + competence.description)
        for ind_code in sorted(competence.indicator_codes):
            indicator = competence.indicator_codes[ind_code]
            set_cell_text(table, row, 1, word_doc.JUSTIFY, ind_code + ' ' + indicator.description)

    # Знать, уметь, владеть
    table.cell(1, 2).merge(table.cell(row, 2))
    add_study_results('knowledge', 'Знать:')
    add_study_results('abilities', 'Уметь:')
    add_study_results('skills', 'Владеть:')

    # Темы и задания
    table.cell(1, 3).merge(table.cell(row, 3))
    for t in course.themes:
        set_cell_text(table, 1, 3, word_doc.CENTER, t['тема'])
    table.cell(1, 4).merge(table.cell(row, 4))
    # for c in course.controls:
    #     set_cell_text(table, 1, 4, word_doc.JUSTIFY, c['содержание'])


def fill_table_7(template: DocxTemplate, context: Dict[str, Any]) -> None:
    """ Заполняем таблицу со ссылками на литературу в разделе 7 """
    if context['lit_images']:
        return  # вместо таблицы будет скан страницы с печатью

    def append_table_7_section(caption, books):
        rows_count = len(table.rows)
        add_table_rows(table, len(books) + 1)  # доп. строка для заголовка
        table.cell(rows_count, 0).merge(table.cell(rows_count, 4))
        set_cell_text(table, rows_count, 0, word_doc.CENTER, caption)
        for i, book in enumerate(books):
            set_cell_text(table, rows_count + i + 1, 0, word_doc.CENTER, str(i + 1))
            set_cell_text(table, rows_count + i + 1, 1, word_doc.CENTER, book['гост'])
            set_cell_text(table, rows_count + i + 1, 2, word_doc.CENTER, book.get('гриф', '—'))
            set_cell_text(table, rows_count + i + 1, 3, word_doc.CENTER, book.get('экз', '—'))
            set_cell_text(table, rows_count + i + 1, 4, word_doc.CENTER, book.get('эбс', '—'))

    table = template.get_docx().tables[10]
    append_table_7_section('Основная литература', context['course'].primary_books)
    append_table_7_section('Дополнительная литература', context['course'].secondary_books)


def remove_extra_table_5(template: DocxTemplate, context: Dict[str, Any]):
    """ Удаляем лишнюю таблицу из раздела 5 """
    exam_table, credit_table = 6, 7
    subject = context['subject']
    control = [s.control for s in subject.semesters.values()]
    control = functools.reduce(lambda x, y: x + y if isinstance(x, list) else set.union(x, y), control)
    if CT_EXAM in control:
        word_doc.remove_table(template, credit_table)
    else:
        word_doc.remove_table(template, exam_table)


def get_images(template: DocxTemplate, subject: Subject, args: argparse.Namespace) -> Dict[str, List[InlineImage]]:
    """
    Поиск картинок, типы которых перечислены в IMAGE_KINDS, например,
    литературы или титульных листов. Папки для поиска указывается в args
    """
    images = {}
    for kind in IMAGE_KINDS:
        try:
            images[kind] = []
            path = vars(args).get(kind + '_dir')
            if path is None:
                continue
            fns = glob.glob(os.path.join(path, '*'))
            pic_subj = {fn: os.path.splitext(os.path.basename(fn))[0] for fn in fns}
            distances = [(fn, levenshtein_d(subject.name, pic_subj[fn])) for fn in pic_subj]
            best = min(distances, key=itemgetter(1))
            images[kind] += glob.glob(os.path.join(path, pic_subj[best[0]]+'*'))
            if best[1]/len(subject.name) < 0.4:
                print(f'Найдены файл(ы) сканов ({kind}): ' + ' '.join(images[kind]))
            elif best[1]/len(subject.name) <= 0.7:
                print(f'Подозрительный скан ({kind}): {best[0]}')
            else:
                print(f'Подходящих сканов не найдено, наименее далекий по имени файл ({kind}): {best[0]}')
                images[kind] = []
            images[kind] = [InlineImage(template, fn, width=Mm(173)) for fn in sorted(images[kind])]
        except OSError:
            print(f'Файл(ы) сканов ({kind}) не найдены!')
            images[kind] = []
    return images


def main(args=None) -> None:
    """ Точка входа """
    if args is None:
        parser = argparse.ArgumentParser()
        parser.add_argument('plan', type=str, help='PLX-файл РУПа')
        parser.add_argument('course', type=str, help='YAML-файл курса')
        parser.add_argument('-t', '--title_dir', type=str, help='Папка со сканами титульных листов')
        parser.add_argument('-l', '--lit_dir', type=str, help='Папка со сканами литературы')
        parser.add_argument('-o', '--output_file', type=str, help='Название выходного файла docx')
        args = parser.parse_args()

    plan = get_plan(args.plan)
    course = get_course(args.course)
    subject = get_subject(plan, course)
    links_before, links_after = plan.find_dependencies(subject, course)
    template = word_doc.get_template('rpd.docx')
    images = get_images(template, subject, args)

    context = {
        'course': course,
        'plan': plan,
        'subject': subject,
        'links_before': links_before,
        'links_after': links_after,
    }
    for kind in IMAGE_KINDS:
        context[kind + '_images'] = images[kind]

    fill_table_1_2(template, context)
    fill_table_2(template, context)
    fill_table_3_1(template, context)
    fill_table_4(template, context)
    fill_table_6_1(template, context)
    fill_table_6_2(template, context)
    fill_table_7(template, context)
    remove_extra_table_5(template, context)

    template.render(context)
    output_file = args.course.replace('.yaml', '.docx')
    if args.output_file:
        output_file = args.output_file
        if not output_file.endswith('.docx'):
            output_file += '.docx'
    try:
        template.save(output_file)
        print(f'Файл {output_file} успешно сохранен')
    except OSError:
        print(f'Ошибка при сохранении файла {output_file}!')


if __name__ == '__main__':
    main()
